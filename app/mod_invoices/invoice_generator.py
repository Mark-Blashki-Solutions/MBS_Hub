#!pip install python-docx

import re
from docx import Document
from datetime import date

def populate_invoice(invoice):
  variables = {
    "{LINE_ITEM_1_QTY}": invoice.line_items[1].quantity,
    "{LINE_ITEM_1_DESC}": invoice.line_items[1].description,
    "{LINE_ITEM_1_GST}": invoice.line_items[1].gst,
    "{LINE_ITEM_1_TOTAL}": invoice.line_items[1].total,
    "{LINE_ITEM_2_QTY}": invoice.line_items[2].quantity,
    "{LINE_ITEM_2_DESC}": invoice.line_items[2].description,
    "{LINE_ITEM_2_GST}": invoice.line_items[2].gst,
    "{LINE_ITEM_2_TOTAL}": invoice.line_items[2].total,
    "{LINE_ITEM_3_QTY}": invoice.line_items[3].quantity,
    "{LINE_ITEM_3_DESC}": invoice.line_items[3].description,
    "{LINE_ITEM_3_GST}": invoice.line_items[3].gst,
    "{LINE_ITEM_3_TOTAL}": invoice.line_items[3].total,
    "{LINE_ITEM_4_QTY}": invoice.line_items[4].quantity,
    "{LINE_ITEM_4_DESC}": invoice.line_items[4].description,
    "{LINE_ITEM_4_GST}": invoice.line_items[4].gst,
    "{LINE_ITEM_4_TOTAL}": invoice.line_items[4].total,
    "{GRAND_TOTAL}": invoice.get_total(),
    "{BUSINESS_NAME}": invoice.client.business_name,
    "{ABN}": invoice.client.abn,
    "{ADDRESS_LINE_1}": invoice.client.address_line_1,
    "{ADDRESS_LINE_2}": invoice.client.address_line_2,
    "{DATE}": date.today().today.strftime("%d %B, %Y"),
    "{INVOICE_NUMBER}": invoice.invoice_number
  }
  file_name = "MBS_INV_" + invoice.invoice_number + "_" + invoice.title + ".docx"
  _fill_template(variables, file_name)

def _fill_template(variables, file_name):
  template_file_path = './../invoice_template.docx'
  
  template_document = Document(template_file_path)
  
  for variable_key, variable_value in variables.items():
    for paragraph in template_document.paragraphs:
      _replace_text_in_paragraph(paragraph, variable_key, variable_value)
    
    for table in template_document.tables:
      for col in table.columns:
        for cell in col.cells:
          for paragraph in cell.paragraphs:
            _replace_text_in_paragraph(paragraph, variable_key, variable_value)
  
  template_document.save(file_name)

def _replace_text_in_paragraph(paragraph, key, value):
    if key in paragraph.text:
        inline = paragraph.runs
        for item in inline:
            if key in item.text:
                item.text = item.text.replace(key, value)
